#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
楽々JC v2.0.0
─────────────────────────────────────────
  機能①  リンク一括設定（旧 LinkForge）
  機能②  PDF一括変換
─────────────────────────────────────────
Mac / Windows 対応
tkinterdnd2 によるドラッグ&ドロップ対応
"""

import os
import sys
import shutil
import threading
import platform
import subprocess
import urllib.request
import urllib.error
from pathlib import Path
from copy import deepcopy
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# ── python-docx ─────────────────────────────────────────────────
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    root = tk.Tk(); root.withdraw()
    messagebox.showerror("エラー",
        "python-docx がインストールされていません。\n\n"
        "ターミナルで以下を実行してください:\n\n"
        "【Mac】  python3 -m pip install python-docx\n"
        "【Win】  python -m pip install python-docx")
    sys.exit(1)

# ── D&D ──────────────────────────────────────────────────────────
_DND_BACKEND = None
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    _DND_BACKEND = "tkdnd"
except ImportError:
    DND_FILES = "DND_Files"

# ════════════════════════════════════════════════════════════════
#  カラーテーマ（全画面共通）
# ════════════════════════════════════════════════════════════════
C = {
    "bg":       "#1B2E4A",
    "surface":  "#1E3759",
    "accent":   "#2A4F8A",
    "primary":  "#E94560",
    "text":     "#FFFFFF",
    "sub":      "#C8DCF5",
    "ok":       "#00D9A5",
    "warn":     "#FFC857",
    "err":      "#FF6B6B",
    "info":     "#48C9F7",
    "input_bg": "#152840",
    "border":   "#3A5A8A",
    "drop_hi":  "#274470",
}

APP_VERSION  = "2.0.0"
GITHUB_USER  = "domechin6248"
GITHUB_REPO  = "LinkForge"
GITHUB_RAW   = f"https://raw.githubusercontent.com/{GITHUB_USER}/{GITHUB_REPO}/main"

HYPERLINK_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/"
    "relationships/hyperlink"
)

LINK_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls",
    ".pptx", ".ppt", ".csv", ".txt",
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff",
    ".zip", ".rtf", ".odt", ".ods", ".odp",
}

PDF_EXTENSIONS = {
    ".doc", ".docx", ".odt", ".rtf",
    ".xls", ".xlsx", ".ods", ".csv",
    ".ppt", ".pptx", ".odp",
    ".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".gif", ".webp",
    ".txt",
}


# ════════════════════════════════════════════════════════════════
#  共通ユーティリティ
# ════════════════════════════════════════════════════════════════

def make_scrollable_frame(parent):
    """スクロール可能なフレームを返す (canvas, inner_frame)"""
    outer = tk.Frame(parent, bg=C["bg"])
    outer.pack(fill=tk.BOTH, expand=True)
    cv = tk.Canvas(outer, bg=C["bg"], highlightthickness=0, bd=0)
    sb = ttk.Scrollbar(outer, orient="vertical", command=cv.yview)
    cv.configure(yscrollcommand=sb.set)
    sb.pack(side=tk.RIGHT, fill=tk.Y)
    cv.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    inner = tk.Frame(cv, bg=C["bg"])
    fid = cv.create_window((0, 0), window=inner, anchor="nw")
    inner.bind("<Configure>", lambda e: cv.configure(scrollregion=cv.bbox("all")))
    cv.bind("<Configure>", lambda e: cv.itemconfig(fid, width=e.width))
    return cv, inner


def make_log_widget(parent):
    """ログテキストウィジェットを返す"""
    log_wrap = tk.Frame(parent, bg=C["accent"], padx=1, pady=1)
    log_wrap.pack(fill=tk.BOTH, expand=True, padx=16, pady=(0, 16))
    tk.Label(log_wrap, text=" 処理ログ",
             font=("Helvetica", 9), bg=C["accent"], fg="#C8DCF5"
             ).pack(anchor="w")
    log_inner = tk.Frame(log_wrap, bg=C["input_bg"])
    log_inner.pack(fill=tk.BOTH, expand=True)
    mono = "Menlo" if platform.system() == "Darwin" else "Consolas"
    txt = tk.Text(
        log_inner, height=8, font=(mono, 9),
        bg=C["input_bg"], fg=C["sub"],
        insertbackground=C["primary"],
        relief=tk.FLAT, bd=0, padx=10, pady=8,
        state=tk.DISABLED, wrap=tk.WORD
    )
    lsb = ttk.Scrollbar(log_inner, command=txt.yview)
    txt.configure(yscrollcommand=lsb.set)
    lsb.pack(side=tk.RIGHT, fill=tk.Y)
    txt.pack(fill=tk.BOTH, expand=True)
    txt.tag_configure("success", foreground=C["ok"])
    txt.tag_configure("error",   foreground=C["err"])
    txt.tag_configure("info",    foreground=C["info"])
    txt.tag_configure("warning", foreground=C["warn"])
    return txt


def log_write(txt_widget, msg, tag=""):
    txt_widget.configure(state=tk.NORMAL)
    if tag:
        txt_widget.insert(tk.END, msg + "\n", tag)
    else:
        txt_widget.insert(tk.END, msg + "\n")
    txt_widget.see(tk.END)
    txt_widget.configure(state=tk.DISABLED)


def section_divider(parent):
    tk.Frame(parent, bg=C["border"], height=1).pack(fill=tk.X, padx=16, pady=(8, 10))


def nav_button(parent, text, command):
    """画面間移動用の小さなボタン"""
    return tk.Button(
        parent, text=text, command=command,
        font=("Helvetica", 9, "bold"),
        bg=C["accent"], fg="#FFFFFF",
        activebackground=C["primary"], activeforeground="white",
        relief=tk.FLAT, bd=0, padx=14, pady=5, cursor="hand2"
    )


# ════════════════════════════════════════════════════════════════
#  DropZone ウィジェット（共通）
# ════════════════════════════════════════════════════════════════

class DropZone(tk.Frame):

    def __init__(self, parent, label_text, hint_text,
                 select_mode="file", file_types=None,
                 allow_multiple=False, **kwargs):
        super().__init__(parent, bg=C["surface"], **kwargs)
        self.select_mode   = select_mode
        self.file_types    = file_types or []
        self.allow_multiple = allow_multiple
        self.selected_paths = []
        self.on_change     = None

        self.configure(
            highlightbackground=C["border"],
            highlightthickness=2,
            highlightcolor=C["primary"]
        )

        self.title_lbl = tk.Label(
            self, text="▸ " + label_text,
            font=("Helvetica", 12, "bold"),
            bg=C["surface"], fg=C["text"]
        )
        self.title_lbl.pack(pady=(12, 2), padx=14, anchor="w")

        self.hint_lbl = tk.Label(
            self, text=hint_text,
            font=("Helvetica", 9),
            bg=C["surface"], fg=C["sub"]
        )
        self.hint_lbl.pack(padx=14, anchor="w")

        pf = tk.Frame(self, bg=C["surface"])
        pf.pack(fill=tk.X, padx=14, pady=(6, 2))

        self.path_var = tk.StringVar()
        self.path_entry = tk.Entry(
            pf, textvariable=self.path_var,
            font=("Helvetica", 9),
            bg=C["input_bg"], fg=C["sub"],
            insertbackground=C["primary"],
            selectbackground=C["primary"],
            relief=tk.FLAT, bd=1
        )
        self.path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 6))
        self.path_entry.insert(0, "ここにパスをペースト...")
        self.path_entry.bind("<FocusIn>",  self._entry_in)
        self.path_entry.bind("<FocusOut>", self._entry_out)
        self.path_entry.bind("<Return>",   self._entry_submit)

        tk.Button(
            pf, text="読込", command=self._entry_submit,
            font=("Helvetica", 9),
            bg=C["accent"], fg="#FFFFFF",
            activebackground=C["primary"], activeforeground="white",
            relief=tk.FLAT, bd=0, padx=10, pady=3, cursor="hand2"
        ).pack(side=tk.LEFT)

        self.info_lbl = tk.Label(
            self, text="",
            font=("Helvetica", 9),
            bg=C["surface"], fg=C["info"],
            justify=tk.LEFT, anchor="w", wraplength=440
        )
        self.info_lbl.pack(padx=14, anchor="w", pady=(2, 0))

        bf = tk.Frame(self, bg=C["surface"])
        bf.pack(padx=14, pady=(6, 12), anchor="w")

        label = "ファイル選択" if select_mode == "file" else "フォルダ選択"
        self.sel_btn = tk.Button(
            bf, text=label, command=self._on_click,
            font=("Helvetica", 9),
            bg=C["accent"], fg="#FFFFFF",
            activebackground=C["primary"], activeforeground="white",
            relief=tk.FLAT, bd=0, padx=12, pady=4, cursor="hand2"
        )
        self.sel_btn.pack(side=tk.LEFT, padx=(0, 6))
        self.clr_btn = None

        for w in [self, self.title_lbl, self.hint_lbl]:
            w.bind("<Button-1>", lambda e: self._on_click())

        if _DND_BACKEND == "tkdnd":
            try:
                self.drop_target_register(DND_FILES)
                self.dnd_bind("<<Drop>>",      self._on_drop)
                self.dnd_bind("<<DragEnter>>", self._on_enter)
                self.dnd_bind("<<DragLeave>>", self._on_leave)
            except Exception:
                pass

    def _on_enter(self, event):
        self.configure(highlightbackground=C["primary"], highlightthickness=3,
                       bg=C["drop_hi"])
        for w in [self.title_lbl, self.hint_lbl, self.info_lbl]:
            w.configure(bg=C["drop_hi"])
        return event.action

    def _on_leave(self, event):
        bg = C["surface"]
        self.configure(bg=bg, highlightthickness=2,
                       highlightbackground=C["ok"] if self.selected_paths else C["border"])
        for w in [self.title_lbl, self.hint_lbl, self.info_lbl]:
            w.configure(bg=bg)
        return event.action

    def _on_drop(self, event):
        self._on_leave(event)
        paths = self._parse_paths(event.data)
        if paths:
            self._set_paths(paths)
        return event.action

    @staticmethod
    def _parse_paths(raw):
        paths, i = [], 0
        while i < len(raw):
            if raw[i] == '{':
                end = raw.index('}', i)
                paths.append(raw[i+1:end]); i = end + 2
            elif raw[i] == ' ':
                i += 1
            else:
                end = raw.find(' ', i)
                if end == -1: end = len(raw)
                paths.append(raw[i:end]); i = end + 1
        return [p for p in paths if p.strip()]

    def _entry_in(self, e=None):
        if self.path_var.get() == "ここにパスをペースト...":
            self.path_entry.delete(0, tk.END)
            self.path_entry.configure(fg=C["text"])

    def _entry_out(self, e=None):
        if not self.path_var.get().strip():
            self.path_entry.configure(fg=C["sub"])
            self.path_entry.insert(0, "ここにパスをペースト...")

    def _entry_submit(self, e=None):
        raw = self.path_var.get().strip()
        if not raw or raw == "ここにパスをペースト...": return
        path = raw.strip("'\"").strip()
        if os.path.exists(path):
            self._set_paths([path])
            self.path_entry.delete(0, tk.END)
            self._entry_out()
        else:
            self.path_entry.configure(fg=C["err"])

    def _on_click(self):
        if self.select_mode == "folder":
            p = filedialog.askdirectory(title="フォルダを選択")
            if p: self._set_paths([p])
        elif self.allow_multiple:
            ps = filedialog.askopenfilenames(
                title="ファイルを選択",
                filetypes=self.file_types or [("All", "*.*")])
            if ps: self._set_paths(list(ps))
        else:
            p = filedialog.askopenfilename(
                title="ファイルを選択",
                filetypes=self.file_types or [("All", "*.*")])
            if p: self._set_paths([p])

    def _set_paths(self, paths):
        if self.select_mode == "folder" and not self.allow_multiple:
            self.selected_paths = [paths[0]]
        elif self.allow_multiple:
            ex = set(self.selected_paths)
            for p in paths:
                if p not in ex:
                    self.selected_paths.append(p); ex.add(p)
        else:
            self.selected_paths = [paths[0]]
        self._update_display()
        if self.on_change: self.on_change()

    def _update_display(self):
        if not self.selected_paths:
            self.info_lbl.config(text="")
            self.configure(highlightbackground=C["border"])
            if self.clr_btn: self.clr_btn.destroy(); self.clr_btn = None
            return
        self.configure(highlightbackground=C["ok"], highlightthickness=2)
        names = [Path(p).name for p in self.selected_paths]
        if len(names) == 1 and self.select_mode == "folder":
            folder = Path(self.selected_paths[0])
            files = [f for f in folder.rglob("*")
                     if f.is_file() and f.suffix.lower() in LINK_EXTENSIONS]
            text = f"✓  {folder.name}/  （{len(files)} ファイル）"
        elif len(names) <= 4:
            text = "\n".join(f"✓  {n}" for n in names)
        else:
            text = "\n".join(f"✓  {n}" for n in names[:3])
            text += f"\n  … 他 {len(names)-3} 件"
        self.info_lbl.config(text=text, fg=C["ok"])
        if not self.clr_btn:
            self.clr_btn = tk.Button(
                self.sel_btn.master, text="クリア", command=self._clear,
                font=("Helvetica", 8),
                bg="#3A1010", fg=C["err"],
                activebackground=C["err"], activeforeground="white",
                relief=tk.FLAT, bd=0, padx=8, pady=3, cursor="hand2"
            )
            self.clr_btn.pack(side=tk.LEFT)

    def _clear(self):
        self.selected_paths = []
        self.info_lbl.config(fg=C["info"])
        self._update_display()
        if self.on_change: self.on_change()


# ════════════════════════════════════════════════════════════════
#  リンク処理コア（既存ロジックをそのまま移植）
# ════════════════════════════════════════════════════════════════

import re as _re
import unicodedata as _ud

def _nfc(s):
    return _ud.normalize("NFC", s)

def _strip_number(name):
    name = name.translate(str.maketrans('０１２３４５６７８９', '0123456789'))
    name = _re.sub(r'^[\d\s\.\-_、。．・()（）【】\[\]「」『』〔〕]+', '', name)
    return name.strip()

def get_file_map(link_dirs):
    fm = {}
    for link_dir in link_dirs:
        base = Path(link_dir)
        for f in sorted(base.rglob("*")):
            if f.is_file() and f.suffix.lower() in LINK_EXTENSIONS:
                rel_str = str(f.relative_to(base.parent)).replace("\\", "/")
                stem_nfc = _nfc(f.stem)
                name_nfc = _nfc(f.name)
                fm[stem_nfc] = rel_str
                fm[name_nfc] = rel_str
                stripped = _strip_number(stem_nfc)
                if stripped:
                    fm[stripped] = rel_str
                    fm[stripped + f.suffix] = rel_str
    return fm

def copy_link_trees(link_dirs, dst_parent):
    for link_dir in link_dirs:
        src = Path(link_dir)
        dst = Path(dst_parent) / src.name
        if dst.exists(): shutil.rmtree(str(dst))
        shutil.copytree(str(src), str(dst))

def iter_all_paragraphs(doc):
    def _tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from _tables(cell.tables)
    yield from doc.paragraphs
    yield from _tables(doc.tables)

def _make_run(text, rpr=None):
    r = OxmlElement("w:r")
    if rpr is not None: r.append(deepcopy(rpr))
    t = OxmlElement("w:t"); t.text = text
    if text != text.strip(): t.set(qn("xml:space"), "preserve")
    r.append(t); return r

def _make_hyperlink(rid, text, rpr=None):
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), rid); hl.set(qn("w:history"), "1")
    r = OxmlElement("w:r")
    rp = OxmlElement("w:rPr")
    rs = OxmlElement("w:rStyle"); rs.set(qn("w:val"), "Hyperlink"); rp.append(rs)
    co = OxmlElement("w:color");  co.set(qn("w:val"), "0563C1");    rp.append(co)
    u  = OxmlElement("w:u");      u.set(qn("w:val"), "single");     rp.append(u)
    skip = {"rStyle", "color", "u"}
    if rpr is not None:
        for c in rpr:
            loc = c.tag.split("}")[-1] if "}" in c.tag else c.tag
            if loc not in skip: rp.append(deepcopy(c))
    r.append(rp)
    t = OxmlElement("w:t"); t.text = text
    if text != text.strip(): t.set(qn("xml:space"), "preserve")
    r.append(t); hl.append(r); return hl

def _find_matches(full, fm):
    keys = sorted(fm.keys(), key=len, reverse=True)
    matches, used = [], set()
    for k in keys:
        pos = 0
        while True:
            idx = full.find(k, pos)
            if idx == -1: break
            end = idx + len(k)
            span = set(range(idx, end))
            if not span & used:
                matches.append((idx, end, fm[k])); used |= span
            pos = idx + 1
    return sorted(matches, key=lambda x: x[0])

def process_paragraph(para, fm, part):
    p = para._p
    runs, pos = [], 0
    for ch in list(p):
        if ch.tag == qn("w:r"):
            te = ch.find(qn("w:t"))
            tx = (te.text or "") if te is not None else ""
            runs.append(dict(elem=ch, text=tx, start=pos,
                             end=pos+len(tx), rpr=ch.find(qn("w:rPr"))))
            pos += len(tx)
    if not runs: return 0
    full = "".join(r["text"] for r in runs)
    ms = _find_matches(_nfc(full), fm)
    if not ms: return 0
    crpr = [None] * len(full)
    for rd in runs:
        for i in range(rd["start"], rd["end"]): crpr[i] = rd["rpr"]
    for rd in runs: p.remove(rd["elem"])
    ppr = p.find(qn("w:pPr"))
    ins = (list(p).index(ppr)+1) if ppr is not None else 0
    elems, prev = [], 0
    for (s, e, rp) in ms:
        if prev < s:
            elems.append(_make_run(full[prev:s], crpr[prev] if prev < len(crpr) else None))
        rid = part.relate_to(rp, HYPERLINK_TYPE, is_external=True)
        elems.append(_make_hyperlink(rid, full[s:e], crpr[s] if s < len(crpr) else None))
        prev = e
    if prev < len(full):
        elems.append(_make_run(full[prev:], crpr[prev] if prev < len(crpr) else None))
    for i, el in enumerate(elems): p.insert(ins+i, el)
    return len(ms)


# ════════════════════════════════════════════════════════════════
#  PDF変換コア
# ════════════════════════════════════════════════════════════════

def _get_libreoffice_path():
    if platform.system() == "Windows":
        for p in [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]:
            if os.path.exists(p): return f'"{p}"'
        return "soffice"
    if platform.system() == "Darwin":
        mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        return mac if os.path.exists(mac) else "soffice"
    return "soffice"

def scan_pdf_targets(paths):
    """フォルダ / ファイルのリストから変換対象を再帰収集
    返り値: [(file_path, root_dir), ...]
    """
    result = []
    for p in paths:
        p = Path(p)
        if p.is_dir():
            root = p
            for f in sorted(p.rglob("*")):
                if f.is_file() and f.suffix.lower() in PDF_EXTENSIONS:
                    result.append((f, root))
        elif p.is_file() and p.suffix.lower() in PDF_EXTENSIONS:
            result.append((p, p.parent))
    return result

def convert_to_pdf(input_path: Path, output_dir: Path, log_cb=None):
    """1ファイルをPDFに変換。成功したらTrue"""
    output_dir.mkdir(parents=True, exist_ok=True)
    lo = _get_libreoffice_path()
    cmd = f'{lo} --headless --convert-to pdf "{input_path}" --outdir "{output_dir}"'
    try:
        result = subprocess.run(
            cmd, shell=True,
            stdout=subprocess.PIPE, stderr=subprocess.PIPE,
            timeout=120
        )
        out_pdf = output_dir / (input_path.stem + ".pdf")
        if out_pdf.exists():
            return True, str(out_pdf)
        else:
            err = result.stderr.decode(errors="ignore")
            return False, err or "出力ファイルが生成されませんでした"
    except subprocess.TimeoutExpired:
        return False, "タイムアウト（120秒）"
    except Exception as e:
        return False, str(e)


# ════════════════════════════════════════════════════════════════
#  共通ヘッダーウィジェット
# ════════════════════════════════════════════════════════════════

def build_header(parent, title, subtitle, show_version=True, on_update=None):
    hdr = tk.Frame(parent, bg=C["accent"], height=56)
    hdr.pack(fill=tk.X); hdr.pack_propagate(False)
    tk.Label(hdr, text=title,
             font=("Helvetica", 18, "bold"),
             bg=C["accent"], fg=C["primary"]
             ).pack(side=tk.LEFT, padx=20, pady=10)
    tk.Label(hdr, text=subtitle,
             font=("Helvetica", 10),
             bg=C["accent"], fg="#C8DCF5"
             ).pack(side=tk.LEFT, pady=10)
    if show_version:
        if on_update:
            tk.Button(hdr, text="🔄", command=on_update,
                      font=("Helvetica", 11),
                      bg=C["accent"], fg="#AACFEE",
                      activebackground=C["primary"], activeforeground="white",
                      relief=tk.FLAT, bd=0, padx=6, cursor="hand2"
                      ).pack(side=tk.RIGHT, pady=10)
        tk.Label(hdr, text=f"v{APP_VERSION}",
                 font=("Helvetica", 8),
                 bg=C["accent"], fg="#AACFEE"
                 ).pack(side=tk.RIGHT, padx=(0, 2), pady=10)
    tk.Frame(parent, bg=C["primary"], height=3).pack(fill=tk.X)
    return hdr


# ════════════════════════════════════════════════════════════════
#  ランチャー画面
# ════════════════════════════════════════════════════════════════

class LauncherFrame(tk.Frame):

    def __init__(self, parent, on_link, on_pdf):
        super().__init__(parent, bg=C["bg"])
        self.on_link = on_link
        self.on_pdf  = on_pdf
        self._build()

    def _build(self):
        build_header(self, "⛓  楽々JC", "自動処理ツール集")

        # ── 中央コンテンツ ──
        center = tk.Frame(self, bg=C["bg"])
        center.pack(expand=True)

        tk.Label(center, text="使用する機能を選んでください",
                 font=("Helvetica", 13),
                 bg=C["bg"], fg=C["text"]
                 ).pack(pady=(40, 32))

        # ── ボタンエリア ──
        btn_area = tk.Frame(center, bg=C["bg"])
        btn_area.pack()

        self._feature_btn(
            btn_area,
            icon="⛓",
            title="リンク一括設定",
            desc="Wordファイルへ\nハイパーリンクを自動挿入",
            command=self.on_link,
            col=0
        )

        self._feature_btn(
            btn_area,
            icon="📄",
            title="PDF一括変換",
            desc="Word / Excel / PowerPoint等を\nPDFに一括変換",
            command=self.on_pdf,
            col=1
        )

        # ── 将来拡張用プレースホルダー ──
        self._placeholder_btn(btn_area, col=2)

        tk.Label(center, text=f"楽々JC  v{APP_VERSION}",
                 font=("Helvetica", 8),
                 bg=C["bg"], fg="#5A7AAA"
                 ).pack(pady=(40, 0))

    def _feature_btn(self, parent, icon, title, desc, command, col):
        frame = tk.Frame(
            parent, bg=C["surface"],
            highlightbackground=C["border"],
            highlightthickness=2,
            cursor="hand2"
        )
        frame.grid(row=0, column=col, padx=14, pady=4, ipadx=6, ipady=6)

        tk.Label(frame, text=icon,
                 font=("Helvetica", 36),
                 bg=C["surface"], fg=C["primary"]
                 ).pack(pady=(24, 6), padx=36)

        tk.Label(frame, text=title,
                 font=("Helvetica", 14, "bold"),
                 bg=C["surface"], fg=C["text"]
                 ).pack()

        tk.Label(frame, text=desc,
                 font=("Helvetica", 9),
                 bg=C["surface"], fg=C["sub"],
                 justify=tk.CENTER
                 ).pack(pady=(4, 20), padx=20)

        # ホバー効果
        def on_enter(e, f=frame):
            f.configure(highlightbackground=C["primary"], highlightthickness=2)
        def on_leave(e, f=frame):
            f.configure(highlightbackground=C["border"], highlightthickness=2)
        def on_click(e):
            command()

        for w in frame.winfo_children() + [frame]:
            w.bind("<Enter>", on_enter)
            w.bind("<Leave>", on_leave)
            w.bind("<Button-1>", on_click)

    def _placeholder_btn(self, parent, col):
        frame = tk.Frame(
            parent, bg=C["surface"],
            highlightbackground=C["border"],
            highlightthickness=1,
        )
        frame.grid(row=0, column=col, padx=14, pady=4, ipadx=6, ipady=6)
        tk.Label(frame, text="＋",
                 font=("Helvetica", 36),
                 bg=C["surface"], fg=C["border"]
                 ).pack(pady=(24, 6), padx=36)
        tk.Label(frame, text="機能追加予定",
                 font=("Helvetica", 12),
                 bg=C["surface"], fg=C["border"]
                 ).pack()
        tk.Label(frame, text="近日公開",
                 font=("Helvetica", 9),
                 bg=C["surface"], fg=C["border"]
                 ).pack(pady=(4, 20), padx=20)


# ════════════════════════════════════════════════════════════════
#  リンク一括設定 画面
# ════════════════════════════════════════════════════════════════

class LinkFrame(tk.Frame):

    def __init__(self, parent, on_back, on_go_pdf, dnd_ok):
        super().__init__(parent, bg=C["bg"])
        self.on_back   = on_back
        self.on_go_pdf = on_go_pdf
        self.dnd_ok    = dnd_ok
        self.folder_name_entries = []
        self._build()

    def _build(self):
        build_header(self, "⛓  リンク一括設定", "Word ハイパーリンク自動挿入",
                     on_update=self._check_update)

        # ── ナビ行 ──
        nav = tk.Frame(self, bg=C["bg"])
        nav.pack(fill=tk.X, padx=16, pady=(8, 0))
        nav_button(nav, "← ホームへ戻る", self.on_back).pack(side=tk.LEFT)
        nav_button(nav, "PDF一括変換へ →", self.on_go_pdf).pack(side=tk.RIGHT)

        # ── D&D状態 ──
        self.dnd_lbl = tk.Label(nav, text="D&D ✓" if self.dnd_ok else "---",
                                 font=("Helvetica", 9),
                                 bg=C["bg"],
                                 fg=C["ok"] if self.dnd_ok else "#AACFEE")
        self.dnd_lbl.pack(side=tk.RIGHT, padx=10)

        # ── スクロールエリア ──
        _, main = make_scrollable_frame(self)
        pad = dict(padx=16, pady=(0, 10))

        self.word_zone = DropZone(
            main, "計画書（Word ファイル）",
            "ドラッグ＆ドロップ、またはクリックして選択（複数可）",
            select_mode="file",
            file_types=[("Word 文書", "*.docx")],
            allow_multiple=True
        )
        self.word_zone.pack(fill=tk.X, **pad)
        self.word_zone.on_change = self._on_word_changed

        self.link_zone = DropZone(
            main, "リンク資料フォルダ",
            "ドラッグ＆ドロップで複数追加可能（何階層でも対応）",
            select_mode="folder", allow_multiple=True
        )
        self.link_zone.pack(fill=tk.X, **pad)
        self.link_zone.on_change = self._check_ready

        self.output_zone = DropZone(
            main, "出力先フォルダ",
            "ドラッグ＆ドロップ、またはクリックして選択",
            select_mode="folder", allow_multiple=False
        )
        self.output_zone.pack(fill=tk.X, **pad)
        self.output_zone.on_change = self._check_ready

        self.fname_outer = tk.Frame(main, bg=C["bg"])

        self.structure_lbl = tk.Label(
            main, text="",
            font=("Helvetica", 9),
            bg=C["bg"], fg=C["sub"],
            justify=tk.LEFT, anchor="w"
        )
        self.structure_lbl.pack(padx=16, anchor="w")

        section_divider(main)

        self.run_btn = tk.Button(
            main,
            text="▶  リンクを作成",
            command=self._run,
            font=("Helvetica", 14, "bold"),
            bg=C["accent"], fg="#AACFEE",
            activebackground=C["primary"], activeforeground="white",
            relief=tk.FLAT, bd=0, padx=30, pady=12,
            cursor="arrow", state=tk.DISABLED
        )
        self.run_btn.pack(pady=(0, 10))

        self.log_txt = make_log_widget(main)
        log_write(self.log_txt,
                  "D&D: 有効 ✓" if self.dnd_ok else "D&D: ダイアログ / パスペーストを利用",
                  "success" if self.dnd_ok else "info")

    # ── リンク機能ロジック（元コードそのまま） ──────────────────

    def _log(self, msg, tag=""):
        self.after(0, lambda m=msg, t=tag: log_write(self.log_txt, m, t))

    def _on_word_changed(self):
        self._rebuild_folder_names()
        self._check_ready()

    def _rebuild_folder_names(self):
        self.fname_outer.pack_forget()
        for w in self.fname_outer.winfo_children(): w.destroy()
        self.folder_name_entries = []
        word_paths = self.word_zone.selected_paths
        if not word_paths:
            self.structure_lbl.configure(text=""); return

        tk.Label(self.fname_outer,
                 text="出力フォルダ名（変更可）",
                 font=("Helvetica", 10, "bold"),
                 bg=C["bg"], fg=C["sub"]
                 ).pack(anchor="w", padx=16, pady=(0, 4))

        for wp in word_paths:
            p = Path(wp)
            if p.name.startswith("~$"): continue
            row = tk.Frame(self.fname_outer, bg=C["bg"])
            row.pack(fill=tk.X, padx=16, pady=2)
            tk.Label(row, text=f"{p.name}  →",
                     font=("Helvetica", 9),
                     bg=C["bg"], fg=C["sub"],
                     anchor="e", width=30
                     ).pack(side=tk.LEFT, padx=(0, 6))
            var = tk.StringVar(value=p.stem)
            tk.Entry(row, textvariable=var,
                     font=("Helvetica", 10),
                     bg=C["input_bg"], fg=C["text"],
                     insertbackground=C["primary"],
                     relief=tk.FLAT, bd=1
                     ).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 6))
            def _reset(v=var, d=p.stem): v.set(d)
            tk.Button(row, text="戻す", command=_reset,
                      font=("Helvetica", 8),
                      bg=C["accent"], fg="#FFFFFF",
                      relief=tk.FLAT, bd=0, padx=6, pady=2, cursor="hand2"
                      ).pack(side=tk.LEFT)
            self.folder_name_entries.append((wp, var))

        self.fname_outer.pack(fill=tk.X, pady=(0, 4), before=self.structure_lbl)
        self._update_structure_label()

    def _update_structure_label(self):
        if not self.folder_name_entries: return
        lines = ["📁 出力フォルダ構成プレビュー:"]
        for wp, var in self.folder_name_entries:
            fn = var.get() or Path(wp).stem
            lines.append(f"  出力先/{fn}/")
            lines.append(f"    ├─ {Path(wp).name}")
            for lp in self.link_zone.selected_paths:
                lines.append(f"    ├─ {Path(lp).name}/")
        self.structure_lbl.configure(text="\n".join(lines))

    def _check_ready(self):
        self._update_structure_label()
        ok = (self.word_zone.selected_paths and
              self.link_zone.selected_paths and
              self.output_zone.selected_paths)
        if ok:
            self.run_btn.configure(state=tk.NORMAL, bg=C["primary"],
                                   fg="white", cursor="hand2")
        else:
            self.run_btn.configure(state=tk.DISABLED, bg=C["accent"],
                                   fg="#AACFEE", cursor="arrow")

    def _run(self):
        self.run_btn.configure(state=tk.DISABLED,
                               text="⏳  処理中...", bg=C["accent"], fg=C["warn"])
        self.log_txt.configure(state=tk.NORMAL)
        self.log_txt.delete("1.0", tk.END)
        self.log_txt.configure(state=tk.DISABLED)
        threading.Thread(target=self._worker, daemon=True).start()

    def _worker(self):
        try:
            link_dirs  = self.link_zone.selected_paths
            output_dir = Path(self.output_zone.selected_paths[0])
            fn_map     = {wp: var.get().strip() or Path(wp).stem
                          for wp, var in self.folder_name_entries}
            fm = get_file_map(link_dirs)
            if not fm:
                self._log("[警告] リンク資料フォルダに対象ファイルがありません。", "warning")
                self.after(0, self._reset_btn); return

            self._log("リンク対象ファイル:", "info")
            seen = set()
            for v in fm.values():
                if v not in seen:
                    self._log(f"  {v}"); seen.add(v)
            self._log("")

            entries = [(wp, fn) for wp, fn in fn_map.items()
                       if not Path(wp).name.startswith("~$")]
            self._log(f"{len(entries)} 件の Word を処理...", "info")
            total = 0
            for wp, cname in entries:
                wpath = Path(wp)
                self._log(f"  {wpath.name}  →  {cname}/")
                try:
                    doc = Document(wpath)
                    out = output_dir / cname
                    out.mkdir(parents=True, exist_ok=True)
                    copy_link_trees(link_dirs, out)
                    self._log(f"    資料フォルダ {len(link_dirs)} 個をコピー完了")
                    cnt = sum(process_paragraph(p, fm, doc.part)
                              for p in iter_all_paragraphs(doc))
                    doc.save(out / wpath.name)
                    if cnt > 0:
                        self._log(f"    {cnt} 箇所リンク挿入", "success")
                    else:
                        self._log("    対象テキストなし", "warning")
                    total += cnt
                except Exception as e:
                    self._log(f"    エラー: {e}", "error")

            self._log("")
            self._log(f"完了！  合計 {total} 箇所のリンクを挿入しました。", "success")
            self._log(f"出力先: {output_dir}", "info")
            self.after(0, lambda: messagebox.showinfo(
                "完了",
                f"処理が完了しました！\n\n合計 {total} 箇所のリンクを挿入\n出力先: {output_dir}"
            ))
        except Exception as e:
            self._log(f"\n予期しないエラー: {e}", "error")
        finally:
            self.after(0, self._reset_btn)

    def _reset_btn(self):
        self.run_btn.configure(text="▶  リンクを作成")
        self._check_ready()

    def _check_update(self):
        def _do():
            try:
                url = GITHUB_RAW + "/version.txt"
                with urllib.request.urlopen(url, timeout=5) as res:
                    latest = res.read().decode().strip()
                if latest == APP_VERSION:
                    self.after(0, lambda: messagebox.showinfo(
                        "アップデート確認", "最新バージョンです！\n現在: v" + APP_VERSION))
                else:
                    lv = latest
                    self.after(0, lambda: messagebox.showinfo(
                        "アップデートあり",
                        f"新しいバージョンがあります。\n現在: v{APP_VERSION}  最新: v{lv}"))
            except Exception as e:
                self.after(0, lambda: messagebox.showerror("エラー", str(e)))
        threading.Thread(target=_do, daemon=True).start()


# ════════════════════════════════════════════════════════════════
#  PDF一括変換 画面
# ════════════════════════════════════════════════════════════════

class PdfFrame(tk.Frame):

    def __init__(self, parent, on_back, on_go_link):
        super().__init__(parent, bg=C["bg"])
        self.on_back    = on_back
        self.on_go_link = on_go_link
        self._cancel_flag = threading.Event()
        self._run_enabled    = False
        self._cancel_enabled = False
        self._build()

    def _build(self):
        build_header(self, "📄  PDF一括変換",
                     "Word / Excel / PowerPoint / 画像 → PDF")

        nav = tk.Frame(self, bg=C["bg"])
        nav.pack(fill=tk.X, padx=16, pady=(8, 0))
        nav_button(nav, "← ホームへ戻る", self.on_back).pack(side=tk.LEFT)
        nav_button(nav, "リンク一括設定へ →", self.on_go_link).pack(side=tk.RIGHT)

        _, main = make_scrollable_frame(self)
        pad = dict(padx=16, pady=(0, 10))

        # 変換対象
        self.src_zone = DropZone(
            main, "変換対象フォルダ / ファイル",
            "複数フォルダをドラッグ&ドロップ（サブフォルダも自動検索）",
            select_mode="folder", allow_multiple=True
        )
        self.src_zone.pack(fill=tk.X, **pad)
        self.src_zone.on_change = self._update_count

        # 出力先
        self.out_zone = DropZone(
            main, "出力先フォルダ",
            "ドラッグ&ドロップ、またはクリックして選択",
            select_mode="folder", allow_multiple=False
        )
        self.out_zone.pack(fill=tk.X, **pad)
        self.out_zone.on_change = self._update_count

        # ファイル件数
        self.count_lbl = tk.Label(
            main, text="",
            font=("Helvetica", 9),
            bg=C["bg"], fg=C["sub"]
        )
        self.count_lbl.pack(padx=16, anchor="w")

        section_divider(main)

        # ボタン行
        btn_row = tk.Frame(main, bg=C["bg"])
        btn_row.pack(pady=(0, 10))

        self.run_btn = tk.Frame(btn_row, bg=C["accent"], cursor="arrow")
        self._run_lbl = tk.Label(self.run_btn,
            text="▶  PDF変換を開始",
            font=("Helvetica", 14, "bold"),
            bg=C["accent"], fg="#AACFEE", padx=30, pady=12)
        self._run_lbl.pack()
        def _run_enter(e):
            if self._run_enabled:
                self.run_btn.configure(bg=C["primary"]); self._run_lbl.configure(bg=C["primary"])
        def _run_leave(e):
            if not self._run_enabled:
                self.run_btn.configure(bg=C["accent"]); self._run_lbl.configure(bg=C["accent"])
        def _run_click(e):
            if self._run_enabled: self._run()
        for _w in (self.run_btn, self._run_lbl):
            _w.bind("<Enter>", _run_enter); _w.bind("<Leave>", _run_leave); _w.bind("<Button-1>", _run_click)
        self.run_btn.pack(side=tk.LEFT, padx=(0, 12))

        self.cancel_btn = tk.Frame(btn_row, bg="#3A1010", cursor="arrow")
        self._cancel_lbl = tk.Label(self.cancel_btn, text="■ 中断",
            font=("Helvetica", 11),
            bg="#3A1010", fg=C["err"], padx=16, pady=12)
        self._cancel_lbl.pack()
        def _can_enter(e):
            if self._cancel_enabled:
                self.cancel_btn.configure(bg=C["err"]); self._cancel_lbl.configure(bg=C["err"], fg="white")
        def _can_leave(e):
            if self._cancel_enabled:
                self.cancel_btn.configure(bg="#3A1010"); self._cancel_lbl.configure(bg="#3A1010", fg=C["err"])
        def _can_click(e):
            if self._cancel_enabled: self._cancel()
        for _w in (self.cancel_btn, self._cancel_lbl):
            _w.bind("<Enter>", _can_enter); _w.bind("<Leave>", _can_leave); _w.bind("<Button-1>", _can_click)
        self.cancel_btn.pack(side=tk.LEFT)

        self.log_txt = make_log_widget(main)
        log_write(self.log_txt, "フォルダをドロップして変換を開始してください", "info")

    def _update_count(self):
        paths = self.src_zone.selected_paths
        if not paths:
            self.count_lbl.configure(text="")
            self._run_enabled = False
            self.run_btn.configure(bg=C["accent"], cursor="arrow")
            self._run_lbl.configure(bg=C["accent"], fg="#AACFEE")
            return
        file_pairs = scan_pdf_targets(paths)
        self.count_lbl.configure(
            text=f"対象ファイル: {len(file_pairs)} 件",
            fg=C["info"] if file_pairs else C["warn"]
        )
        if file_pairs:
            self._run_enabled = True
            self.run_btn.configure(bg=C["primary"], cursor="hand2")
            self._run_lbl.configure(bg=C["primary"], fg="#FFFFFF")
        else:
            self._run_enabled = False
            self.run_btn.configure(bg=C["accent"], cursor="arrow")
            self._run_lbl.configure(bg=C["accent"], fg="#AACFEE")

    def _log(self, msg, tag=""):
        self.after(0, lambda m=msg, t=tag: log_write(self.log_txt, m, t))

    def _cancel(self):
        self._cancel_flag.set()
        self._log("中断リクエストを送信しました...", "warning")

    def _run(self):
        self._cancel_flag.clear()
        self._run_enabled = False
        self.run_btn.configure(bg=C["accent"], cursor="arrow")
        self._run_lbl.configure(bg=C["accent"], fg=C["warn"], text="⏳  変換中...")
        self._cancel_enabled = True
        self.cancel_btn.configure(bg="#3A1010", cursor="hand2")
        self._cancel_lbl.configure(bg="#3A1010", fg=C["err"])
        self.log_txt.configure(state=tk.NORMAL)
        self.log_txt.delete("1.0", tk.END)
        self.log_txt.configure(state=tk.DISABLED)
        threading.Thread(target=self._worker, daemon=True).start()

    def _worker(self):
        try:
            src_paths = self.src_zone.selected_paths
            file_pairs = scan_pdf_targets(src_paths)
            total = len(file_pairs)
            self._log(f"対象ファイル {total} 件を変換します...", "info")

            out_paths = self.out_zone.selected_paths
            base_out = Path(out_paths[0]) if out_paths else None

            success, fail = 0, 0
            for i, (f, root) in enumerate(file_pairs, 1):
                if self._cancel_flag.is_set():
                    self._log(f"中断しました（{i-1}/{total} 件処理済）", "warning")
                    break

                self._log(f"[{i}/{total}] {f.name}")

                if base_out:
                    try:
                        rel = f.parent.relative_to(root)
                        out_dir = base_out / root.name / rel
                    except ValueError:
                        out_dir = base_out / root.name
                else:
                    out_dir = f.parent

                ok, detail = convert_to_pdf(f, out_dir)
                if ok:
                    success += 1
                    self._log(f"  ✓ → {f.stem}.pdf", "success")
                else:
                    fail += 1
                    self._log(f"  ✕ {detail}", "error")

            self._log("")
            self._log(f"完了！  成功 {success} 件 ／ 失敗 {fail} 件", "success")
            self.after(0, lambda: messagebox.showinfo(
                "PDF変換完了",
                f"変換が完了しました！\n\n成功: {success} 件\n失敗: {fail} 件"
            ))
        except Exception as e:
            self._log(f"予期しないエラー: {e}", "error")
        finally:
            self.after(0, self._reset_btns)

    def _reset_btns(self):
        self._cancel_enabled = False
        self.cancel_btn.configure(bg="#3A1010", cursor="arrow")
        self._cancel_lbl.configure(bg="#3A1010", fg=C["err"])
        self._run_lbl.configure(text="▶  PDF変換を開始")
        self._update_count()


# ════════════════════════════════════════════════════════════════
#  メインアプリ（画面切り替えコントローラ）
# ════════════════════════════════════════════════════════════════

class RakurakuJC:

    def __init__(self):
        if _DND_BACKEND == "tkdnd":
            try:
                self.root = TkinterDnD.Tk()
            except Exception:
                self.root = tk.Tk()
        else:
            self.root = tk.Tk()

        self.root.title(f"楽々JC  v{APP_VERSION}")
        self.root.geometry("680x800")
        self.root.minsize(560, 680)
        self.root.configure(bg=C["bg"])

        if platform.system() == "Darwin":
            self.root.lift()
            self.root.attributes("-topmost", True)
            self.root.after(100, lambda: self.root.attributes("-topmost", False))

        self._current = None
        self._show_launcher()

    # ── 画面管理 ──────────────────────────────────────────────────

    def _clear(self):
        if self._current:
            self._current.destroy()
            self._current = None

    def _show_launcher(self):
        self._clear()
        f = LauncherFrame(
            self.root,
            on_link=self._show_link,
            on_pdf=self._show_pdf
        )
        f.pack(fill=tk.BOTH, expand=True)
        self._current = f

    def _show_link(self):
        self._clear()
        f = LinkFrame(
            self.root,
            on_back=self._show_launcher,
            on_go_pdf=self._show_pdf,
            dnd_ok=(_DND_BACKEND == "tkdnd")
        )
        f.pack(fill=tk.BOTH, expand=True)
        self._current = f

    def _show_pdf(self):
        self._clear()
        f = PdfFrame(
            self.root,
            on_back=self._show_launcher,
            on_go_link=self._show_link
        )
        f.pack(fill=tk.BOTH, expand=True)
        self._current = f

    def run(self):
        self.root.mainloop()


# ════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    RakurakuJC().run()
